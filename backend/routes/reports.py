import json
import os
from flask import request, jsonify, send_file
from utils.db import get_db
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import base64


def create_reports_table():
    """Create the reports table if it doesn't exist."""
    conn = get_db()
    if not conn:
        return False
    try:
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS reports (
                id INT AUTO_INCREMENT PRIMARY KEY,
                session_id VARCHAR(64) NOT NULL,
                report_type VARCHAR(50),
                report_data JSON,
                file_path VARCHAR(512),
                generated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                INDEX idx_session_id (session_id)
            )
        """)
        conn.commit()
        cursor.close()
        conn.close()
        return True
    except Exception as e:
        print(f"Error creating reports table: {e}")
        conn.close()
        return False


def generate_appendix_data(data, instrument_type):
    """Generate appendix data for reports."""
    appendix = {
        'instrument_details': [],
        'metadata': {
            'instrument_type': instrument_type,
            'generated_at': datetime.now().isoformat(),
            'total_instruments': len(data) if data else 0
        }
    }
    
    if not data or not isinstance(data, list):
        return appendix
    
    for idx, item in enumerate(data):
        instrument_detail = {
            'index': idx + 1,
            'name': item.get('Instrument') or item.get('BondName') or item.get('TBillName') or f'Instrument {idx + 1}',
            'face_value': item.get('FaceValue') or item.get('Amount') or item.get('Principal') or 0,
            'rate': item.get('Rate') or item.get('InterestRate') or item.get('CouponRate') or item.get('DiscountRate') or 0,
            'term': item.get('Term') or item.get('YearsToMaturity') or item.get('DaysToMaturity') or 0,
            'issue_date': item.get('IssueDate') or item.get('Issue Date') or 'N/A',
            'maturity_date': item.get('MaturityDate') or item.get('Maturity') or 'N/A',
            'currency': item.get('Currency') or 'USD',
            'calculated_value': item.get('CalculatedValue') or 0,
            'difference': item.get('Difference') or 0,
            'yield': item.get('Yield') or 0
        }
        appendix['instrument_details'].append(instrument_detail)
    
    return appendix


def generate_report_excel(session_id, report_data, instrument_type):
    """Generate Excel report with cover, summary, data, and appendix sheets."""
    output_dir = 'reports'
    os.makedirs(output_dir, exist_ok=True)
    
    wb = openpyxl.Workbook()
    
    title_font = Font(name='Arial', size=16, bold=True, color='0B2044')
    header_font = Font(name='Arial', size=12, bold=True, color='0B2044')
    normal_font = Font(name='Arial', size=10)
    
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')
    wrap_align = Alignment(horizontal='left', vertical='top', wrap_text=True)
    
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Sheet 1: Cover Page
    ws_cover = wb.active
    ws_cover.title = 'Cover'
    
    ws_cover['A1'] = 'DURA CAPITAL (PRIVATE) LIMITED'
    ws_cover['A1'].font = Font(name='Arial', size=20, bold=True, color='0B2044')
    ws_cover['A1'].alignment = center_align
    
    ws_cover['A3'] = 'PORTFOLIO VALUATION REPORT'
    ws_cover['A3'].font = Font(name='Arial', size=18, bold=True, color='1E88E5')
    ws_cover['A3'].alignment = center_align
    
    ws_cover['A5'] = f'Instrument Type: {instrument_type.replace("-", " ").upper()}'
    ws_cover['A5'].font = header_font
    ws_cover['A5'].alignment = left_align
    
    ws_cover['A6'] = f'Session ID: {session_id}'
    ws_cover['A6'].font = normal_font
    ws_cover['A6'].alignment = left_align
    
    ws_cover['A7'] = f'Report Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}'
    ws_cover['A7'].font = normal_font
    ws_cover['A7'].alignment = left_align
    
    ws_cover['A9'] = 'CONFIDENTIAL'
    ws_cover['A9'].font = Font(name='Arial', size=14, bold=True, color='c62828')
    ws_cover['A9'].alignment = center_align
    
    ws_cover.column_dimensions['A'].width = 50
    
    # Sheet 2: Summary
    ws_summary = wb.create_sheet('Summary')
    
    summary_data = report_data.get('summary', {})
    
    ws_summary['A1'] = 'PORTFOLIO SUMMARY'
    ws_summary['A1'].font = title_font
    ws_summary['A1'].alignment = center_align
    ws_summary.merge_cells('A1:B1')
    
    ws_summary['A3'] = 'Metric'
    ws_summary['B3'] = 'Value'
    ws_summary['A3'].font = header_font
    ws_summary['B3'].font = header_font
    ws_summary['A3'].alignment = center_align
    ws_summary['B3'].alignment = center_align
    
    metrics = [
        ('Total Portfolio Value', summary_data.get('total_value', 0)),
        ('Number of Instruments', summary_data.get('instrument_count', 0)),
        ('Average Rate (%)', summary_data.get('avg_rate', 0)),
        ('Total Interest Earned', summary_data.get('total_interest', 0)),
        ('Total Principal', summary_data.get('total_principal', 0)),
    ]
    
    for idx, (label, value) in enumerate(metrics, start=4):
        ws_summary[f'A{idx}'] = label
        ws_summary[f'B{idx}'] = value
        ws_summary[f'A{idx}'].font = normal_font
        ws_summary[f'B{idx}'].font = normal_font
        ws_summary[f'A{idx}'].alignment = left_align
        ws_summary[f'B{idx}'].alignment = right_align
    
    ws_summary[f'A{len(metrics)+4}'] = 'Generated at'
    ws_summary[f'B{len(metrics)+4}'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    ws_summary[f'A{len(metrics)+4}'].font = normal_font
    ws_summary[f'B{len(metrics)+4}'].font = normal_font
    
    ws_summary.column_dimensions['A'].width = 30
    ws_summary.column_dimensions['B'].width = 20
    
    # Sheet 3: Data
    ws_data = wb.create_sheet('Data')
    data = report_data.get('data', [])
    if data and len(data) > 0:
        headers = list(data[0].keys())
        ws_data.append(headers)
        
        for col_idx, header in enumerate(headers, start=1):
            cell = ws_data.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.alignment = center_align
            cell.border = thin_border
        
        for row_idx, row in enumerate(data, start=2):
            for col_idx, header in enumerate(headers, start=1):
                cell = ws_data.cell(row=row_idx, column=col_idx)
                cell.value = row.get(header, '')
                cell.font = normal_font
                cell.alignment = wrap_align
                cell.border = thin_border
        
        # Auto-fit columns
        for col in ws_data.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws_data.column_dimensions[column].width = adjusted_width
    
    # Sheet 4: Appendix
    ws_appendix = wb.create_sheet('Appendix')
    appendix = generate_appendix_data(data, instrument_type)
    
    ws_appendix['A1'] = 'APPENDIX: DETAILED INSTRUMENT DATA'
    ws_appendix['A1'].font = title_font
    ws_appendix['A1'].alignment = center_align
    ws_appendix.merge_cells('A1:J1')
    
    appendix_headers = ['Instrument Name', 'Face Value', 'Rate (%)', 'Term', 'Issue Date', 'Maturity Date', 'Currency', 'Calculated Value', 'Difference', 'Yield (%)']
    ws_appendix.append([''])
    ws_appendix.append(appendix_headers)
    
    for col_idx, header in enumerate(appendix_headers, start=1):
        cell = ws_appendix.cell(row=3, column=col_idx)
        cell.font = header_font
        cell.alignment = center_align
        cell.border = thin_border
    
    for row_idx, detail in enumerate(appendix['instrument_details'], start=4):
        values = [
            detail['name'],
            detail['face_value'],
            detail['rate'],
            detail['term'],
            detail['issue_date'],
            detail['maturity_date'],
            detail['currency'],
            detail['calculated_value'],
            detail['difference'],
            detail['yield']
        ]
        for col_idx, value in enumerate(values, start=1):
            cell = ws_appendix.cell(row=row_idx, column=col_idx)
            cell.value = value
            cell.font = normal_font
            cell.alignment = left_align
            cell.border = thin_border
    
    ws_appendix.column_dimensions['A'].width = 25
    ws_appendix.column_dimensions['B'].width = 15
    ws_appendix.column_dimensions['C'].width = 12
    ws_appendix.column_dimensions['D'].width = 12
    ws_appendix.column_dimensions['E'].width = 15
    ws_appendix.column_dimensions['F'].width = 15
    ws_appendix.column_dimensions['G'].width = 12
    ws_appendix.column_dimensions['H'].width = 18
    ws_appendix.column_dimensions['I'].width = 15
    ws_appendix.column_dimensions['J'].width = 12
    
    # Set page setup for A4
    for sheet in wb.worksheets:
        sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
        sheet.page_setup.orientation = sheet.ORIENTATION_PORTRAIT
        sheet.page_setup.fitToPage = True
        sheet.page_setup.fitToHeight = 0
        sheet.page_setup.fitToWidth = 1
        sheet.page_margins.left = 0.5
        sheet.page_margins.right = 0.5
        sheet.page_margins.top = 0.5
        sheet.page_margins.bottom = 0.5
    
    file_name = f"report_{session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    file_path = os.path.join(output_dir, file_name)
    wb.save(file_path)
    
    return file_path


def save_report(session_id, report_type, report_data, file_path):
    """Save report metadata to database."""
    conn = get_db()
    if not conn:
        return None
    try:
        cursor = conn.cursor()
        cursor.execute(
            """INSERT INTO reports (session_id, report_type, report_data, file_path) 
               VALUES (%s, %s, %s, %s)""",
            (session_id, report_type, json.dumps(report_data), file_path)
        )
        conn.commit()
        report_id = cursor.lastrowid
        cursor.close()
        conn.close()
        return report_id
    except Exception as e:
        print(f"Error saving report: {e}")
        conn.close()
        return None


def generate_report_word(session_id, report_data, instrument_type, fred_filters=None, yield_curve_data=None, chart_image_base64=None, logo_base64=None, background_base64=None):
    """Generate Word document with cover page, TOC, and full report content."""
    output_dir = 'reports'
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    
    # Set document margins (narrow for professional look)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Define styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # ===== PAGE 1: COVER PAGE =====
    # Cover page content (no page break before - this is the first page)
    
    # Add logo if available
    if logo_base64:
        try:
            if logo_base64.startswith('data:image'):
                logo_base64 = logo_base64.split(',')[1]
            logo_data = base64.b64decode(logo_base64)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_logo:
                temp_logo.write(logo_data)
                temp_logo.flush()
                # Add logo at top left
                logo_paragraph = doc.add_paragraph()
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(temp_logo.name, width=Inches(2))
                os.unlink(temp_logo.name)
        except Exception as e:
            print(f"Error adding logo: {e}")
    
    # Add background image if available (as a watermark-style image)
    if background_base64:
        try:
            if background_base64.startswith('data:image'):
                background_base64 = background_base64.split(',')[1]
            bg_data = base64.b64decode(background_base64)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_bg:
                temp_bg.write(bg_data)
                temp_bg.flush()
                # Add background image on the right side
                bg_paragraph = doc.add_paragraph()
                bg_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                bg_run = bg_paragraph.add_run()
                bg_run.add_picture(temp_bg.name, width=Inches(3))
                os.unlink(temp_bg.name)
        except Exception as e:
            print(f"Error adding background: {e}")
    
    cover_paragraph = doc.add_paragraph()
    cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company name
    company_run = cover_paragraph.add_run('DURA CAPITAL (PRIVATE) LIMITED')
    company_run.font.size = Pt(20)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(11, 32, 68)
    
    doc.add_paragraph().add_run()
    
    # Report title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run('VALUATION ASSESSMENT REPORT')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 136, 229)
    
    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()
    
    # Session name
    session_paragraph = doc.add_paragraph()
    session_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    session_run = session_paragraph.add_run(f'Session: {report_data.get("session", "Current Session")}')
    session_run.font.size = Pt(16)
    session_run.font.bold = True
    
    # Instrument type
    inst_paragraph = doc.add_paragraph()
    inst_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_run = inst_paragraph.add_run(f'Instrument: {instrument_type.replace("-", " ").upper()}')
    inst_run.font.size = Pt(14)
    
    # Date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    date_run.font.size = Pt(12)
    
    doc.add_paragraph().add_run()
    doc.add_paragraph().add_run()
    
    # Confidential
    conf_paragraph = doc.add_paragraph()
    conf_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_paragraph.add_run('CONFIDENTIAL')
    conf_run.font.size = Pt(14)
    conf_run.font.bold = True
    conf_run.font.color.rgb = RGBColor(198, 40, 40)
    
    # Add page break after cover page
    doc.add_page_break()
    
    # ===== PAGE 2: TABLE OF CONTENTS =====
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        ('Introduction', '3'),
        ('Executive Summary', '4'),
        ('Methodology', '5'),
        ('Market Inputs', '6'),
        ('Results', '7'),
        ('Yield Curve', '8'),
        ('Conclusion', '9'),
        ('Appendix: Detailed Instrument Data', '10'),
        ('Reference', '11')
    ]
    
    for item, page in toc_items:
        toc_para = doc.add_paragraph()
        toc_para.add_run(item)
        toc_para.add_run(f'\t{page}')
        toc_para.paragraph_format.tab_stops.add_tab_stop(Cm(15))
    
    # Add page break after TOC
    doc.add_page_break()
    
    # ===== PAGE 3: INTRODUCTION =====
    doc.add_heading('Introduction', level=1)
    
    intro_para = doc.add_paragraph()
    intro_para.add_run(f'Dura Capital (Private) Limited was contracted to provide a fair valuation assessment report of the following {instrument_type} instruments as at {report_data.get("valuationDate", datetime.now().strftime("%Y-%m-%d"))}.')
    
    # Add bullet list matching HTML
    intro_list = doc.add_paragraph(style='List Bullet')
    intro_list.add_run(f'{instrument_type} instruments')
    intro_list = doc.add_paragraph(style='List Bullet')
    intro_list.add_run(f'Valuation as at {report_data.get("valuationDate", datetime.now().strftime("%Y-%m-%d"))}')
    
    data = report_data.get('data', [])
    if data:
        intro_list = doc.add_paragraph(style='List Bullet')
        intro_list.add_run(f'{len(data)} individual instruments assessed')
    
    # ===== PAGE 4: EXECUTIVE SUMMARY =====
    doc.add_page_break()
    doc.add_heading('Executive Summary', level=1)
    
    # Calculate summary metrics from actual data
    total_value = 0
    total_interest = 0
    rates = []
    
    if data:
        for item in data:
            total_value += float(item.get('Total Value', item.get('Calculated Value', item.get('FaceValue', 0))))
            total_interest += float(item.get('Total Interest', item.get('Interest Earned', 0)))
            rate = float(item.get('Avg Rate', item.get('Rate', item.get('InterestRate', 0))))
            if rate > 0:
                rates.append(rate)
    
    avg_rate = sum(rates) / len(rates) if rates else 0
    
    # Add methodology based on instrument type (matching HTML)
    inst_type = instrument_type.lower()
    if 'money' in inst_type:
        methodology = 'Money Market Instruments: Short-term debt instruments valued using discounted cash flow methodology.'
    elif 'bond' in inst_type:
        methodology = 'Corporate Bonds: Fixed income securities valued using present value of future cash flows.'
    elif 'tbill' in inst_type or 't-bill' in inst_type:
        methodology = 'Treasury Bills: Short-term government securities valued using discount yield methodology.'
    else:
        methodology = 'General fixed income valuation methodology.'
    
    # Create executive summary box (matching HTML styling)
    summary_box = doc.add_paragraph()
    summary_box_format = summary_box.paragraph_format
    summary_box_format.left_indent = Inches(0.25)
    summary_box_format.right_indent = Inches(0.25)
    summary_box_format.space_before = Pt(12)
    summary_box_format.space_after = Pt(12)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Key Findings:').bold = True
    
    summary_list = [
        f'Total Portfolio Value: ${total_value:,.2f}',
        f'Number of Instruments: {len(data)}',
        f'Average Rate: {avg_rate:.2f}%',
        f'Valuation Date: {report_data.get("valuationDate", datetime.now().strftime("%Y-%m-%d"))}'
    ]
    
    for item in summary_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_paragraph().add_run('\nValuation Approach: ')
    doc.add_paragraph(methodology)
    
    # ===== PAGE 5: METHODOLOGY =====
    doc.add_page_break()
    doc.add_heading('Methodology', level=1)
    
    # Create methodology box (matching HTML styling)
    meth_box = doc.add_paragraph()
    meth_box_format = meth_box.paragraph_format
    meth_box_format.left_indent = Inches(0.25)
    meth_box_format.right_indent = Inches(0.25)
    meth_box_format.space_before = Pt(12)
    meth_box_format.space_after = Pt(12)
    
    doc.add_paragraph(methodology)
    
    doc.add_paragraph().add_run('\nFormulas:').bold = True
    
    if 'money' in inst_type:
        formula = 'Fair value = F / (1 + r·t/365) where F = Face value, r = annualized interest rate, t = days to maturity.'
        assumptions = 'Simple interest convention (365 days/year). Weighted average rate = Σ (Rate × Amount) / Σ Amount.'
    elif 'bond' in inst_type:
        formula = 'Fair value = Σ C/(1+y)^t + FV/(1+y)^n where C = annual coupon payment, y = yield to maturity, FV = face value, n = years to maturity.'
        assumptions = 'Coupon payments are annualized. Duration = Σ (t × PV(C_t)) / Price.'
    elif 'tbill' in inst_type or 't-bill' in inst_type:
        formula = 'Discount amount = Face value × (Discount rate/100) × (Days to maturity/360). Effective yield = (Face value / Price − 1) × (365 / Days to maturity) × 100.'
        assumptions = 'Bank discount basis (360 days/year) for discount rate; bond equivalent yield uses 365 days.'
    else:
        formula = 'Present value of expected future cash flows discounted at appropriate market rates.'
        assumptions = 'Standard market conventions applied.'
    
    # Formula box (matching HTML styling)
    formula_box = doc.add_paragraph()
    formula_box_format = formula_box.paragraph_format
    formula_box_format.left_indent = Inches(0.25)
    formula_box_format.right_indent = Inches(0.25)
    formula_box_format.space_before = Pt(12)
    formula_box_format.space_after = Pt(12)
    
    formula_para = doc.add_paragraph()
    formula_run = formula_para.add_run(formula)
    formula_run.font.name = 'Courier New'
    formula_run.font.size = Pt(10)
    
    doc.add_paragraph().add_run('\nAssumptions: ').bold = True
    doc.add_paragraph(assumptions)
    
    # ===== PAGE 6: MARKET INPUTS =====
    doc.add_page_break()
    doc.add_heading('Market Inputs', level=1)
    
    if fred_filters:
        market_para = doc.add_paragraph()
        market_para.add_run(f'Rates sourced from FRED for {report_data.get("valuationDate", datetime.now().strftime("%Y-%m-%d"))}. Filters used:')
        market_para.add_run('\n\n')
        market_para.add_run(f'Country: {fred_filters.get("country", "US")}')
        market_para.add_run('\n')
        market_para.add_run(f'Currency: {fred_filters.get("currency", "USD")}')
        market_para.add_run('\n')
        market_para.add_run(f'Maturity: {fred_filters.get("maturity", "1Y")}')
    else:
        doc.add_paragraph('Rates sourced from FRED for valuation.')
    
    # ===== PAGE 7: RESULTS =====
    doc.add_page_break()
    doc.add_heading('Results', level=1)
    
    # Create results table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True
    
    results_data = [
        ('Total Portfolio Value', f'${total_value:,.2f}'),
        ('Number of Instruments', str(len(data))),
        ('Average Rate', f'{avg_rate:.2f}%'),
        ('Total Interest Earned', f'${total_interest:,.2f}'),
        ('Valuation Date', report_data.get('valuationDate', datetime.now().strftime("%Y-%m-%d")))
    ]
    
    for metric, value in results_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    # ===== PAGE 8: YIELD CURVE =====
    doc.add_page_break()
    doc.add_heading('Yield Curve', level=1)
    
    doc.add_paragraph('The following yield curve was used as a benchmark for valuation, sourced from FRED.')
    
    # Add chart image if available (matching HTML structure)
    if chart_image_base64:
        try:
            # Remove data URL prefix if present
            if chart_image_base64.startswith('data:image'):
                chart_image_base64 = chart_image_base64.split(',')[1]
            
            image_data = base64.b64decode(chart_image_base64)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_img:
                temp_img.write(image_data)
                temp_img.flush()
                # Add chart with styling matching HTML
                chart_paragraph = doc.add_paragraph()
                chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_run = chart_paragraph.add_run()
                chart_run.add_picture(temp_img.name, width=Inches(6))
                os.unlink(temp_img.name)
            
            # Add caption matching HTML
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f'FRED Yield Curve – {instrument_type} ({fred_filters.get("country", "US")} / {fred_filters.get("currency", "USD")})')
            caption_run.font.size = Pt(10)
            caption_run.font.italic = True
        except Exception as e:
            print(f"Error adding yield curve chart: {e}")
            doc.add_paragraph('Yield curve chart could not be embedded.')
    else:
        doc.add_paragraph('Yield curve chart not available.')
    
    # ===== PAGE 9: CONCLUSION =====
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run(f'The valuation assessment is in accordance with IFRS 13 fair value measurement principles as at {report_data.get("valuationDate", datetime.now().strftime("%Y-%m-%d"))}.')
    
    # ===== PAGE 10: APPENDIX =====
    doc.add_page_break()
    doc.add_heading('Appendix: Detailed Instrument Data', level=1)
    
    if data:
        # Create instrument data table (matching HTML structure)
        inst_table = doc.add_table(rows=1, cols=6)
        inst_table.style = 'Table Grid'
        
        headers = ['Instrument Name', 'BB Ticker', 'Face Value ($)', 'Rate (%)', 'Term (Yrs)', 'Valuation Date']
        hdr_cells = inst_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            # Style header to match HTML (dark blue background)
            hdr_cells[i].background_color = RGBColor(11, 32, 68)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        for item in data:
            row_cells = inst_table.add_row().cells
            row_cells[0].text = item.get('Instrument', item.get('BondName', item.get('TBillName', 'N/A')))
            row_cells[1].text = item.get('BBTicker', item.get('Ticker', item.get('Security', 'N/A')))
            row_cells[2].text = str(item.get('FaceValue', item.get('Amount', item.get('Principal', 0))))
            row_cells[3].text = str(item.get('Rate', item.get('InterestRate', item.get('CouponRate', 0))))
            row_cells[4].text = str(item.get('Term', item.get('YearsToMaturity', 0)))
            row_cells[5].text = report_data.get('valuationDate', datetime.now().strftime("%Y-%m-%d"))
        
        # Add FRED yield curve data if available (matching HTML structure)
        if yield_curve_data and len(yield_curve_data) > 0:
            doc.add_paragraph().add_run()
            doc.add_heading('FRED Yield Curve Data', level=2)
            
            fred_para = doc.add_paragraph()
            fred_para.add_run(f'Country: {fred_filters.get("country", "US")} | Currency: {fred_filters.get("currency", "USD")} | Maturity: {fred_filters.get("maturity", "1Y")}')
            
            fred_table = doc.add_table(rows=1, cols=3)
            fred_table.style = 'Table Grid'
            
            fred_headers = ['Maturity Label', 'Term (Yr)', 'Rate (%)']
            fred_hdr_cells = fred_table.rows[0].cells
            for i, header in enumerate(fred_headers):
                fred_hdr_cells[i].text = header
                fred_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                # Style header to match HTML (dark blue background)
                fred_hdr_cells[i].background_color = RGBColor(26, 58, 110)
                fred_hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            for point in yield_curve_data:
                row_cells = fred_table.add_row().cells
                row_cells[0].text = point.get('maturityLabel', point.get('maturity', ''))
                row_cells[1].text = str(point.get('maturity', 0))
                row_cells[2].text = str(point.get('rate', 0))
    
    # ===== PAGE 11: REFERENCE =====
    doc.add_page_break()
    doc.add_heading('Reference', level=1)
    
    ref_list = [
        'FRED – Federal Reserve Economic Data',
        'IFRS 13: Fair Value Measurement',
        'IFRS 9: Financial Instruments'
    ]
    
    for ref in ref_list:
        doc.add_paragraph(ref, style='List Bullet')
    
    # Footer
    doc.add_paragraph().add_run()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f'© {datetime.now().year} Dura Capital (Private) Limited. Report generated {datetime.now().strftime("%B %d, %Y at %H:%M")}.')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    file_name = f"report_{session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(output_dir, file_name)
    doc.save(file_path)
    
    return file_path


def reports_routes(app):
    """Register all report routes."""
    
    # Create table on module load
    create_reports_table()
    
    @app.route('/api/report/generate', methods=['POST', 'OPTIONS'])
    def generate_report_endpoint():
        if request.method == 'OPTIONS':
            return '', 200
        
        payload = request.get_json() or {}
        session_id = payload.get('session_id')
        report_type = payload.get('report_type', 'portfolio')
        instrument_type = payload.get('instrument_type', 'money-market')
        report_data = payload.get('report_data', {})
        
        if not session_id:
            return jsonify({'success': False, 'message': 'Session ID is required'}), 400
        
        file_path = generate_report_excel(session_id, report_data, instrument_type)
        report_id = save_report(session_id, report_type, report_data, file_path)
        
        if report_id:
            return jsonify({
                'success': True,
                'data': {
                    'report_id': report_id,
                    'file_path': file_path,
                    'report_type': report_type
                }
            })
        else:
            return jsonify({'success': False, 'message': 'Failed to save report'}), 500
    
    @app.route('/api/report/appendix', methods=['POST', 'OPTIONS'])
    def generate_appendix_endpoint():
        if request.method == 'OPTIONS':
            return '', 200
        
        payload = request.get_json() or {}
        data = payload.get('data', [])
        instrument_type = payload.get('instrument_type', 'money-market')
        
        appendix = generate_appendix_data(data, instrument_type)
        
        return jsonify({
            'success': True,
            'data': appendix
        })
    
    @app.route('/api/report/<int:report_id>/download', methods=['GET', 'OPTIONS'])
    def download_report(report_id):
        if request.method == 'OPTIONS':
            return '', 200
        
        conn = get_db()
        if not conn:
            return jsonify({'success': False, 'message': 'Database error'}), 500
        
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM reports WHERE id = %s", (report_id,))
            report = cursor.fetchone()
            cursor.close()
            conn.close()
            
            if not report:
                return jsonify({'success': False, 'message': 'Report not found'}), 404
            
            file_path = report['file_path']
            
            if not os.path.exists(file_path):
                return jsonify({'success': False, 'message': 'Report file not found'}), 404
            
            return send_file(file_path, as_attachment=True, download_name=os.path.basename(file_path))
            
        except Exception as e:
            conn.close()
            return jsonify({'success': False, 'message': str(e)}), 500
    
    @app.route('/api/report/generate-word', methods=['POST', 'OPTIONS'])
    def generate_word_report_endpoint():
        if request.method == 'OPTIONS':
            return '', 200
        
        payload = request.get_json() or {}
        session_id = payload.get('session_id')
        report_data = payload.get('report_data', {})
        instrument_type = payload.get('instrument_type', 'money-market')
        fred_filters = payload.get('fred_filters')
        yield_curve_data = payload.get('yield_curve_data')
        chart_image_base64 = payload.get('chart_image_base64')
        logo_base64 = payload.get('logo_base64')
        background_base64 = payload.get('background_base64')
        
        if not session_id:
            return jsonify({'success': False, 'message': 'Session ID is required'}), 400
        
        try:
            file_path = generate_report_word(
                session_id,
                report_data,
                instrument_type,
                fred_filters,
                yield_curve_data,
                chart_image_base64,
                logo_base64,
                background_base64
            )
            
            if file_path and os.path.exists(file_path):
                return jsonify({
                    'success': True,
                    'data': {
                        'file_path': file_path,
                        'file_name': os.path.basename(file_path)
                    }
                })
            else:
                return jsonify({'success': False, 'message': 'Failed to generate Word document'}), 500
                
        except Exception as e:
            print(f"Error generating Word report: {e}")
            return jsonify({'success': False, 'message': str(e)}), 500
    
    @app.route('/api/report/download-word/<session_id>', methods=['GET', 'OPTIONS'])
    def download_word_report(session_id):
        if request.method == 'OPTIONS':
            return '', 200
        
        try:
            # Find the most recent Word report for this session
            output_dir = 'reports'
            if not os.path.exists(output_dir):
                return jsonify({'success': False, 'message': 'Reports directory not found'}), 404
            
            # List all .docx files for this session
            word_files = [f for f in os.listdir(output_dir) if f.startswith(f'report_{session_id}_') and f.endswith('.docx')]
            
            if not word_files:
                return jsonify({'success': False, 'message': 'No Word report found for this session'}), 404
            
            # Get the most recent file
            word_files.sort(reverse=True)
            latest_file = word_files[0]
            file_path = os.path.join(output_dir, latest_file)
            
            return send_file(file_path, as_attachment=True, download_name=latest_file)
            
        except Exception as e:
            print(f"Error downloading Word report: {e}")
            return jsonify({'success': False, 'message': str(e)}), 500

    @app.route('/api/report/session/<session_id>', methods=['GET', 'OPTIONS'])
    def get_session_reports(session_id):
        if request.method == 'OPTIONS':
            return '', 200
        
        conn = get_db()
        if not conn:
            return jsonify({'success': False, 'message': 'Database error'}), 500
        
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM reports WHERE session_id = %s ORDER BY generated_at DESC", (session_id,))
            rows = cursor.fetchall()
            cursor.close()
            conn.close()
            
            reports = []
            for row in rows:
                reports.append({
                    'id': row['id'],
                    'session_id': row['session_id'],
                    'report_type': row['report_type'],
                    'file_path': row['file_path'],
                    'generated_at': row['generated_at'].isoformat() if row['generated_at'] else None
                })
            
            return jsonify({'success': True, 'data': reports})
            
        except Exception as e:
            conn.close()
            return jsonify({'success': False, 'message': str(e)}), 500